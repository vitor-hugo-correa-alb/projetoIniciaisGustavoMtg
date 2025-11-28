# NOTE: Este é o generate_word.py com as últimas melhorias (heurísticas de títulos, modelo_final com 3 títulos, replace_bar_placeholder, etc.).
import logging
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from copy import deepcopy
import zipfile
import tempfile
import os

# logger para este módulo
logger = logging.getLogger(__name__)

# Try to import Composer from docxcompose (best option to preserve images/relationships)
try:
    from docxcompose.composer import Composer  # type: ignore
    logger.info("docxcompose.Composer disponível — imagens serão preservadas quando anexar modelos.")
except Exception as e:
    Composer = None
    logger.info("docxcompose.Composer não disponível: %s. Uso fallback para anexação (pode perder imagens).", e)

def int_to_roman(num: int) -> str:
    vals = [
        (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'), (100, 'C'),
        (90, 'XC'), (50, 'L'), (40, 'XL'), (10, 'X'),
        (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
    ]
    roman = ''
    for val, symbol in vals:
        while num >= val:
            roman += symbol
            num -= val
    return roman

# Placeholders that should be forced bold when substituted
CAMPOS_NEGRITO = {
    "{{NOME_RECLAMANTE}}",
    "{{PROFISSAO_RECLAMANTE}}",
    "{{NOME_MAE_RECLAMANTE}}"
}

def _copy_run_formatting(src_run, dst_run):
    try:
        if src_run.font.name:
            dst_run.font.name = src_run.font.name
            try:
                dst_run._element.rPr.rFonts.set(qn('w:eastAsia'), src_run.font.name)
            except Exception:
                pass
    except Exception:
        pass
    try:
        if src_run.font.size:
            dst_run.font.size = src_run.font.size
    except Exception:
        pass
    try:
        dst_run.bold = src_run.bold
    except Exception:
        pass
    try:
        dst_run.italic = src_run.italic
    except Exception:
        pass
    try:
        dst_run.underline = src_run.underline
    except Exception:
        pass

def replace_placeholders_in_paragraph_preserve_runs(paragraph, substitutions, campos_negrito):
    runs = list(paragraph.runs)
    if not runs:
        return

    run_texts = [r.text for r in runs]
    full_text = ''.join(run_texts)

    if not any(ph in full_text for ph in substitutions.keys()):
        return

    cursor = 0
    segments = []

    def _find_run_index_and_offset(pos):
        acc = 0
        for idx, txt in enumerate(run_texts):
            if acc + len(txt) > pos:
                return idx, pos - acc
            acc += len(txt)
        return len(run_texts) - 1, max(0, len(run_texts[-1]))

    while cursor < len(full_text):
        matched = False
        for placeholder, valor in substitutions.items():
            if full_text.startswith(placeholder, cursor):
                matched = True
                valor_text = "" if valor is None else str(valor)
                start_pos = cursor
                end_pos = cursor + len(placeholder)

                start_run_idx, start_offset = _find_run_index_and_offset(start_pos)
                end_run_idx, end_offset = _find_run_index_and_offset(end_pos - 1)

                before_text = run_texts[start_run_idx][:start_offset]
                after_text = run_texts[end_run_idx][end_offset:]

                for i in range(0, start_run_idx):
                    segments.append( (run_texts[i], runs[i]) )

                if before_text:
                    segments.append( (before_text, runs[start_run_idx]) )

                segments.append( (valor_text, runs[start_run_idx], placeholder in campos_negrito) )

                for i in range(end_run_idx + 1, len(runs)):
                    segments.append( (run_texts[i], runs[i]) )

                for r in list(paragraph.runs):
                    try:
                        r._element.getparent().remove(r._element)
                    except Exception:
                        pass

                for seg in segments:
                    if len(seg) == 2:
                        text, src_run = seg
                        if text == "":
                            continue
                        new_run = paragraph.add_run(text)
                        _copy_run_formatting(src_run, new_run)
                    else:
                        text, src_run, force_bold = seg
                        new_run = paragraph.add_run(text)
                        try:
                            _copy_run_formatting(src_run, new_run)
                        except Exception:
                            pass
                        try:
                            if not new_run.font.name:
                                new_run.font.name = "Garamond"
                                try:
                                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Garamond')
                                except Exception:
                                    pass
                            if not new_run.font.size:
                                new_run.font.size = Pt(12)
                        except Exception:
                            pass
                        if force_bold:
                            new_run.bold = True

                replace_placeholders_in_paragraph_preserve_runs(paragraph, substitutions, campos_negrito)
                return
        if not matched:
            cursor += 1

def replace_placeholders_in_table(table, substitutions, campos_negrito):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_placeholders_in_paragraph_preserve_runs(p, substitutions, campos_negrito)
            for tbl in cell.tables:
                replace_placeholders_in_table(tbl, substitutions, campos_negrito)

def replace_placeholders_in_doc(doc, substitutions, campos_negrito):
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph_preserve_runs(p, substitutions, campos_negrito)
    for t in doc.tables:
        replace_placeholders_in_table(t, substitutions, campos_negrito)
    try:
        for section in doc.sections:
            header = section.header
            for p in header.paragraphs:
                replace_placeholders_in_paragraph_preserve_runs(p, substitutions, campos_negrito)
            for t in header.tables:
                replace_placeholders_in_table(t, substitutions, campos_negrito)
            footer = section.footer
            for p in footer.paragraphs:
                replace_placeholders_in_paragraph_preserve_runs(p, substitutions, campos_negrito)
            for t in footer.tables:
                replace_placeholders_in_table(t, substitutions, campos_negrito)
    except Exception:
        pass

def _docx_has_media(docx_path: str) -> bool:
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            for name in z.namelist():
                if name.startswith('word/media/'):
                    return True
    except Exception:
        return False
    return False

# regex para detectar prefixo romano seguido por separador (tab, hífen ou –/—)
_ROMAN_PREFIX_RE = re.compile(r'^\s*[IVXLCDM]+\s*[\t\-\u2013\u2014]\s*', flags=re.I)

def replace_bar_placeholder(text: str) -> str:
    if text is None:
        return text
    return text.replace("{{BARRA}}", "/")

def _extract_title_text(paragraph):
    """Return paragraph text with roman prefix stripped (if present)."""
    text = (paragraph.text or "").strip()
    return _ROMAN_PREFIX_RE.sub('', text).strip() or text

def _is_heading_style(paragraph):
    try:
        style = paragraph.style
        if style and getattr(style, 'name', None):
            name = style.name.lower()
            if 'heading' in name or 'title' in name:
                return True
    except Exception:
        pass
    return False

def _paragraph_has_bold(paragraph):
    try:
        for r in paragraph.runs:
            if getattr(r, 'bold', False):
                return True
    except Exception:
        pass
    return False

def _find_title_paragraph_indices(doc: Document, max_count: int = 3):
    """
    Heurística para localizar índices de parágrafos que funcionam como títulos:
    1) parágrafos com estilo contendo 'heading' ou 'title'
    2) parágrafos com runs bold
    3) parágrafos curtos seguidos por parágrafo não-vazio
    4) fallback: primeiros parágrafos não vazios
    Retorna lista de índices (ordenados na ordem que aparecem no documento), até max_count.
    """
    indices = []
    total = len(doc.paragraphs)

    # helper to add index if valid and not duplicate
    def add_idx(i):
        if i is None or i < 0 or i >= total:
            return
        if i not in indices:
            indices.append(i)

    # 1) heading style
    for i, p in enumerate(doc.paragraphs):
        if len(indices) >= max_count:
            break
        if not (p.text or "").strip():
            continue
        if _is_heading_style(p):
            add_idx(i)

    # 2) bold paragraphs
    if len(indices) < max_count:
        for i, p in enumerate(doc.paragraphs):
            if len(indices) >= max_count:
                break
            if i in indices:
                continue
            if not (p.text or "").strip():
                continue
            if _paragraph_has_bold(p):
                add_idx(i)

    # 3) short paragraph followed by non-empty paragraph (title+content pattern)
    if len(indices) < max_count:
        for i in range(len(doc.paragraphs)-1):
            if len(indices) >= max_count:
                break
            if i in indices:
                continue
            p = doc.paragraphs[i]
            nxt = doc.paragraphs[i+1]
            if not (p.text or "").strip():
                continue
            if not (nxt.text or "").strip():
                continue
            if len((p.text or "").strip()) <= 200:
                add_idx(i)

    # 4) fallback: first non-empty paragraphs
    if len(indices) < max_count:
        for i, p in enumerate(doc.paragraphs):
            if len(indices) >= max_count:
                break
            if i in indices:
                continue
            if (p.text or "").strip():
                add_idx(i)

    indices.sort()
    return indices

def _apply_sequential_titles_to_doc(doc: Document, start_idx: int, count: int = 3, prefix_spaces: str = "                  "):
    """
    Localiza até 'count' parágrafos que parecem títulos e substitui seus textos por
    títulos sequenciais começando em start_idx (em números romanos).
    Retorna next_idx = start_idx + number_of_titles_applied
    """
    title_indices = _find_title_paragraph_indices(doc, max_count=count)
    applied = 0
    idx = start_idx

    # Replace existing title paragraphs
    for ti in title_indices:
        if applied >= count:
            break
        try:
            p = doc.paragraphs[ti]
            original_text = _extract_title_text(p)
            # remove runs
            for r in list(p.runs):
                try:
                    r._element.getparent().remove(r._element)
                except Exception:
                    pass
            # insert new numbered title
            title_str = f"{prefix_spaces}{int_to_roman(idx)} - {replace_bar_placeholder(original_text)}"
            new_run = p.add_run(title_str)
            new_run.bold = True
            try:
                new_run.font.name = 'Garamond'
                new_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Garamond')
                new_run.font.size = Pt(12)
            except Exception:
                pass
            applied += 1
            idx += 1
        except Exception as e:
            logger.exception("Erro ao aplicar título sequencial no parágrafo %s: %s", ti, e)
            # continue trying others

    # If not enough titles found, insert missing titles after the last applied title (or at end)
    if applied < count:
        # find insertion point: after last applied index, otherwise end of document
        insert_pos = title_indices[-1] + 1 if title_indices else len(doc.paragraphs)
        for _ in range(count - applied):
            try:
                p_new = doc.add_paragraph("")  # appended at end
                # attempt to move it to insert_pos by inserting before doc.paragraphs[insert_pos] if possible
                try:
                    target = doc.paragraphs[insert_pos]._element
                    # remove runs from p_new then add run
                    for r in list(p_new.runs):
                        try:
                            r._element.getparent().remove(r._element)
                        except Exception:
                            pass
                    title_str = f"{prefix_spaces}{int_to_roman(idx)}- "
                    new_run = p_new.add_run(title_str)
                    new_run.bold = True
                    try:
                        new_run.font.name = 'Garamond'
                        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Garamond')
                        new_run.font.size = Pt(12)
                    except Exception:
                        pass
                    # move element before target
                    p_new._element.addprevious(target)
                except Exception:
                    # fallback: leave appended at end
                    for r in list(p_new.runs):
                        try:
                            r._element.getparent().remove(r._element)
                        except Exception:
                            pass
                    title_str = f"{prefix_spaces}{int_to_roman(idx)}- "
                    new_run = p_new.add_run(title_str)
                    new_run.bold = True
                    try:
                        new_run.font.name = 'Garamond'
                        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Garamond')
                        new_run.font.size = Pt(12)
                    except Exception:
                        pass
                applied += 1
                idx += 1
                insert_pos += 1
            except Exception as e:
                logger.exception("Erro ao inserir título adicional: %s", e)
                idx += 1
                inserted = False

    return start_idx + applied

def gerar_documento(caminho_modelo, campos, pedidos_ordenados, numeracao_inicial=6, ai_response=None):
    logger.info("[generate_word] inicio gerar_documento; modelo_base=%s; pedidos_ordenados=%s", caminho_modelo, pedidos_ordenados)
    # Open base template
    try:
        doc = Document(caminho_modelo)
    except Exception as e:
        logger.exception("Erro ao abrir modelo_base: %s", caminho_modelo)
        raise

    substitutions = {
        "{{NOME_RECLAMANTE}}": campos.get("Nome reclamante", ""),
        "{{PROFISSAO_RECLAMANTE}}": campos.get("Profissao / Cargo", ""),
        "{{DATA_NASCIMENTO_RECLAMANTE}}": campos.get("Data de nascimento", ""),
        "{{NOME_MAE_RECLAMANTE}}": campos.get("Nome da mae", ""),
        "{{NUMERO_PIS_RECLAMANTE}}": campos.get("Número do pis", ""),
        "{{NUMERO_CTPS_RECLAMANTE}}": campos.get("Número da ctps", ""),
        "{{NUMERO_RG_RECLAMANTE}}": campos.get("Número rg", ""),
        "{{NUMERO_CPF_RECLAMANTE}}": campos.get("Número do cpf", ""),
        "{{RUA_DO_RECLAMANTE}}": campos.get("Rua do reclamante", ""),
        "{{NUMERO_CASA_RECLAMANTE_E_COMPLEMENTO}}": campos.get("Número da casa do reclamante e complemento", ""),
        "{{BAIRRO_RECLAMANTE}}": campos.get("Bairro reclamante", ""),
        "{{CEP_RECLAMANTE}}": campos.get("Cep reclamante", ""),
        "{{NOME_RECLAMADA}}": campos.get("Nome da reclamada", ""),
        "{{EMPRESA_PROCESSADA}}": campos.get("Empresa processada", ""),
        "{{NUMERO_CNPJ_RECLAMADA}}": campos.get("Numero de cnpj da reclamada", ""),
        "{{ENDERECO_RECLAMADA}}": campos.get("Endereço reclamada", ""),
        "{{COMPLEMENTO_RECLAMADA}}": campos.get("Complemento reclamada", ""),
        "{{BAIRRO_RECLAMADA}}": campos.get("Bairro reclamada", ""),
        "{{CEP_RECLAMADA}}": campos.get("Cep reclamada", "")
    }

    replace_placeholders_in_doc(doc, substitutions, CAMPOS_NEGRITO)

    if not pedidos_ordenados:
        logger.info("[generate_word] nenhum pedido informado; retornando apenas template.")
        return doc

    # Normalize and filter pedidos_ordenados to existing files
    pedidos_validos = []
    for nome_modelo, caminho in pedidos_ordenados:
        if not caminho:
            logger.warning("[generate_word] caminho vazio para modelo %s; pulando.", nome_modelo)
            continue
        if not os.path.isfile(caminho):
            logger.warning("[generate_word] arquivo modelo não existe: %s; pulando %s.", caminho, nome_modelo)
            continue
        pedidos_validos.append((nome_modelo, caminho))
    if not pedidos_validos:
        logger.info("[generate_word] após filtragem não há modelos válidos; retornando template.")
        return doc

    # caminho do modelo_final (mesmo diretório do modelo_base)
    base_templates_dir = os.path.dirname(caminho_modelo)
    final_model_path = os.path.join(base_templates_dir, "modelo_base_final.docx")
    final_exists = os.path.isfile(final_model_path)

    # If Composer is available, try to use it (best option to preserve images/relationships)
    if Composer is not None:
        composer = None
        try:
            composer = Composer(doc)
        except Exception as e:
            logger.exception("Erro inicializando Composer: %s. Faremos fallback.", e)
            composer = None

        if composer is not None:
            composer_failed = False
            # Try appending using composer. For titles we append a tiny Document with the title before the model.
            for idx, (nome_modelo, caminho_modelo_pedido) in enumerate(pedidos_validos, numeracao_inicial):
                titulo = f"                  {int_to_roman(idx)} - {replace_bar_placeholder(nome_modelo)}"
                logger.info("[generate_word] Composer: anexando %s -> %s", nome_modelo, caminho_modelo_pedido)
                try:
                    # create small Document for title
                    title_doc = Document()
                    p_title = title_doc.add_paragraph()
                    run_title = p_title.add_run(titulo)
                    run_title.bold = True
                    try:
                        run_title.font.name = 'Garamond'
                        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Garamond')
                        run_title.font.size = Pt(12)
                    except Exception:
                        pass

                    # append title doc then the source doc
                    composer.append(title_doc)
                    src = Document(caminho_modelo_pedido)
                    composer.append(src)
                except Exception as e:
                    logger.exception("Composer falhou ao anexar %s: %s", nome_modelo, e)
                    composer_failed = True
                    break

            if not composer_failed:
                # Append final model if exists
                if final_exists:
                    try:
                        final_src = Document(final_model_path)
                        # apply sequential titles to the first 3 title-like paragraphs
                        next_idx = numeracao_inicial + len(pedidos_validos)
                        _apply_sequential_titles_to_doc(final_src, next_idx, count=3)
                        composer.append(final_src)
                        logger.info("[generate_word] Composer: anexado modelo_final %s (com 3 títulos atualizados)", final_model_path)
                    except Exception as e:
                        logger.exception("Composer falhou ao anexar modelo_final: %s", e)
                        composer_failed = True

                if not composer_failed:
                    tmpf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmpf.close()
                    try:
                        composer.save(tmpf.name)
                        final_doc = Document(tmpf.name)
                    finally:
                        try:
                            os.remove(tmpf.name)
                        except Exception:
                            pass
                    logger.info("[generate_word] composer finalizado e documento retornado.")
                    return final_doc
                else:
                    logger.warning("[generate_word] composer falhou durante append final; faremos fallback para anexar restantes via deepcopy.")
                    # If composer failed, fall through to fallback logic below.

    # Composer is not available (or failed). Check whether any model has media; if so, raise an informative error.
    models_with_media = [c for _, c in pedidos_validos if _docx_has_media(c)]
    # Also check if final model has media (it will be appended too)
    if final_exists and _docx_has_media(final_model_path):
        models_with_media.append(final_model_path)

    if models_with_media and Composer is None:
        logger.error("Alguns modelos contêm imagens/mídia e Composer não está disponível: %s", models_with_media)
        raise RuntimeError(
            "Alguns modelos contêm imagens/mídia. Para preservar imagens ao concatenar documentos, instale "
            "a dependência 'docxcompose' no ambiente de build/execution: pip install docxcompose\n"
            "Ou empacote seu executável a partir de um ambiente onde docxcompose esteja instalado."
        )

    # Fallback when no media (or composer failed): append by deep-copying XML elements (preserves most formatting for text/tables)
    for idx, (nome_modelo, caminho_modelo_pedido) in enumerate(pedidos_validos, numeracao_inicial):
        logger.info("[generate_word] Fallback: anexando por deepcopy %s -> %s", nome_modelo, caminho_modelo_pedido)
        titulo = f"                  {int_to_roman(idx)} - {replace_bar_placeholder(nome_modelo)}"
        p_title = doc.add_paragraph()
        run = p_title.add_run(titulo)
        run.bold = True
        try:
            run.font.name = 'Garamond'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Garamond')
            run.font.size = Pt(12)
        except Exception:
            pass

        try:
            src = Document(caminho_modelo_pedido)
            for element in src.element.body:
                doc.element.body.append(deepcopy(element))
        except Exception as e:
            logger.exception("deepcopy falhou para %s: %s; tentando fallback run-a-run.", nome_modelo, e)
            # Last-resort fallback: copy paragraph by paragraph preserving run formatting where possible
            try:
                src = Document(caminho_modelo_pedido)
                for paragraph in src.paragraphs:
                    p = doc.add_paragraph()
                    try:
                        p.style = paragraph.style
                    except Exception:
                        pass
                    for run in paragraph.runs:
                        r = p.add_run(run.text)
                        _copy_run_formatting(run, r)
            except Exception as e2:
                logger.exception("fallback run-a-run também falhou para %s: %s; pula modelo.", nome_modelo, e2)
                pass

        doc.add_paragraph("")

    # Após anexar todos os pedidos, anexamos o modelo_final (se existir) e atualizamos os seus 3 títulos
    if final_exists:
        logger.info("[generate_word] Fallback: anexando modelo_final por deepcopy %s", final_model_path)
        try:
            final_src = Document(final_model_path)
            next_idx = numeracao_inicial + len(pedidos_validos)
            # apply sequential titles to the first 3 title-like paragraphs (in-place)
            _apply_sequential_titles_to_doc(final_src, next_idx, count=3)
            try:
                for element in final_src.element.body:
                    doc.element.body.append(deepcopy(element))
            except Exception as e:
                logger.exception("deepcopy falhou para modelo_final %s: %s; tentando fallback run-a-run.", final_model_path, e)
                try:
                    for paragraph in final_src.paragraphs:
                        p = doc.add_paragraph()
                        try:
                            p.style = paragraph.style
                        except Exception:
                            pass
                        for run in paragraph.runs:
                            r = p.add_run(run.text)
                            _copy_run_formatting(run, r)
                except Exception as e2:
                    logger.exception("fallback run-a-run também falhou para modelo_final %s: %s; pula modelo.", final_model_path, e2)
                    pass

            doc.add_paragraph("")
            logger.info("[generate_word] modelo_final anexado por fallback com títulos atualizados.")
        except Exception as e:
            logger.exception("Erro ao anexar modelo_final: %s", e)
            # não abortamos; retornamos o documento já gerado sem o final se houver erro

    logger.info("[generate_word] anexação por fallback concluída; retornando documento.")
    return doc

def salvar_documento(doc, caminho_destino):
    """
    Salva o docx no caminho_destino de forma mais robusta:
    - Salva primeiro em arquivo temporário na mesma pasta.
    - Move/replace o arquivo temporário para o destino (os.replace) para evitar problemas de escrita direta.
    - Em caso de erro, levanta RuntimeError com mensagem detalhada.
    """
    logger.info("[generate_word] Salvando documento em: %s", caminho_destino)
    if not caminho_destino:
        logger.error("[generate_word] Caminho destino vazio ao salvar documento.")
        raise RuntimeError("Caminho destino vazio ao salvar documento.")

    destino_dir = os.path.dirname(caminho_destino) or os.getcwd()
    if not os.path.isdir(destino_dir):
        # tenta criar diretório se não existir
        try:
            os.makedirs(destino_dir, exist_ok=True)
        except Exception as e:
            logger.exception("Não foi possível criar diretório '%s': %s", destino_dir, e)
            raise RuntimeError(f"Não foi possível criar diretório '{destino_dir}': {e}")

    # Cria um arquivo temporário no mesmo diretório
    fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=destino_dir)
    os.close(fd)
    try:
        # tenta salvar no temporário
        try:
            doc.save(tmp_path)
        except Exception as e:
            logger.exception("Erro ao salvar documento temporário em %s: %s", tmp_path, e)
            raise RuntimeError(f"Erro ao salvar documento temporário em {tmp_path}: {e}")

        # Tenta substituir o destino de forma atômica
        try:
            os.replace(tmp_path, caminho_destino)
        except PermissionError:
            # tenta remover arquivo destino e renomear
            try:
                if os.path.exists(caminho_destino):
                    os.remove(caminho_destino)
                os.replace(tmp_path, caminho_destino)
            except Exception as e:
                try:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)
                except Exception:
                    pass
                logger.exception("Permissão negada ao substituir '%s': %s", caminho_destino, e)
                raise RuntimeError(f"Permissão negada ao substituir '{caminho_destino}': {e}")
        except Exception as e:
            try:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass
            logger.exception("Erro ao mover arquivo temporário para destino '%s': %s", caminho_destino, e)
            raise RuntimeError(f"Erro ao mover arquivo temporário para destino '{caminho_destino}': {e}")

    finally:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass